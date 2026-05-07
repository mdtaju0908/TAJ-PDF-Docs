import os
import re
from html import escape
from typing import List, Optional

from fastapi import UploadFile, HTTPException
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from docx import Document

from app.utils.file_handler import create_temp_file, save_upload_files


def _normalize_markdown(markdown_text: str) -> str:
    text = markdown_text.replace("\r\n", "\n")
    # Remove fenced code block markers while keeping code content.
    text = re.sub(r"^```[\w-]*\s*$", "", text, flags=re.MULTILINE)
    return text


def _markdown_to_pdf(markdown_text: str) -> str:
    out_path, out_id = create_temp_file(".pdf")
    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    x_margin = 0.75 * inch
    y_margin = 0.75 * inch
    y = height - y_margin
    line_height = 14

    for raw_line in _normalize_markdown(markdown_text).splitlines():
        line = raw_line.rstrip() or " "
        # Keep simple wrapping so long lines do not overflow.
        chunks = [line[i : i + 110] for i in range(0, len(line), 110)] or [" "]
        for chunk in chunks:
            if y < y_margin + line_height:
                c.showPage()
                y = height - y_margin
            c.drawString(x_margin, y, chunk)
            y -= line_height

    c.save()
    return out_id


def _markdown_to_docx(markdown_text: str) -> str:
    out_path, out_id = create_temp_file(".docx")
    doc = Document()

    for raw_line in _normalize_markdown(markdown_text).splitlines():
        stripped = raw_line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            doc.add_heading(heading_match.group(2).strip(), level=level)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(stripped)

    doc.save(out_path)
    return out_id


def _markdown_to_doc(markdown_text: str) -> str:
    # Legacy .doc generation via HTML content for broad editor compatibility.
    out_path, out_id = create_temp_file(".doc")
    lines: List[str] = []
    for raw_line in _normalize_markdown(markdown_text).splitlines():
        stripped = raw_line.strip()
        if not stripped:
            lines.append("<p>&nbsp;</p>")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            lines.append(f"<h{level}>{escape(heading_match.group(2).strip())}</h{level}>")
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            lines.append(f"<p>&bull; {escape(stripped[2:].strip())}</p>")
            continue

        lines.append(f"<p>{escape(stripped)}</p>")

    html_doc = (
        "<html><head><meta charset='utf-8'></head><body>"
        + "".join(lines)
        + "</body></html>"
    )
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html_doc)
    return out_id


async def _run(files: List[UploadFile], output_format: str) -> str:
    paths: List[str] = []
    try:
        paths = await save_upload_files(files, allowed_exts=[".md"])
        with open(paths[0], "r", encoding="utf-8", errors="ignore") as f:
            markdown_text = f.read()

        if output_format == "pdf":
            return _markdown_to_pdf(markdown_text)
        if output_format == "doc":
            return _markdown_to_doc(markdown_text)
        if output_format == "docx":
            return _markdown_to_docx(markdown_text)
        raise HTTPException(status_code=400, detail="Unsupported markdown output format")
    finally:
        for p in paths:
            try:
                os.remove(p)
            except Exception:
                pass


async def run_pdf(files: List[UploadFile], options: Optional[dict] = None) -> str:
    return await _run(files, "pdf")


async def run_doc(files: List[UploadFile], options: Optional[dict] = None) -> str:
    return await _run(files, "doc")


async def run_docx(files: List[UploadFile], options: Optional[dict] = None) -> str:
    return await _run(files, "docx")


async def run(files: List[UploadFile], options: Optional[dict] = None) -> str:
    opts = options or {}
    output_format = str(opts.get("output_format", "pdf")).strip().lower()
    if output_format not in {"pdf", "doc", "docx"}:
        raise HTTPException(status_code=400, detail="Unsupported output_format. Use pdf, doc or docx.")
    return await _run(files, output_format)
